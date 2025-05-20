"""
Chunker service implementation.
"""
import asyncio
from typing import List, Dict, Any, Optional, Tuple, Union
import os
import mimetypes
from io import BytesIO

from chatbot.chunker_service.core.config import settings
from chatbot.chunker_service.core.errors import ChunkingError, ValidationError, ResourceNotFoundError
from chatbot.chunker_service.core.logging import setup_logging
from chatbot.chunker_service.chunking import get_chunking_strategy
from chatbot.chunker_service.storage import get_storage_adapter
from chatbot.chunker_service.services.embedding_client import EmbeddingClient
from chatbot.chunker_service.services.document_tracker import DocumentTracker
from chatbot.chunker_service.db.models import ProcessedDocument


logger = setup_logging(__name__)


class ChunkerService:
    """Service for chunking documents and storing embeddings."""

    def __init__(
        self,
        embedding_client: Optional[EmbeddingClient] = None,
        document_tracker: Optional[DocumentTracker] = None
    ):
        """Initialize the chunker service.

        Args:
            embedding_client: Optional embedding client
            document_tracker: Optional document tracker
        """
        # Initialize clients
        self.embedding_client = embedding_client or EmbeddingClient()
        self.document_tracker = document_tracker or DocumentTracker()

        logger.info("Initialized ChunkerService")

    async def chunk_document(
        self,
        document_path: str,
        collection_name: Optional[str] = None,
        chunking_strategy: Optional[str] = None,
        chunking_params: Optional[Dict[str, Any]] = None,
        embedding_model: Optional[str] = None,
        storage_type: Optional[str] = None,
        storage_params: Optional[Dict[str, Any]] = None,
        document_metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Chunk a document and store embeddings.

        Args:
            document_path: Path to the document in storage
            collection_name: Name of the collection to store embeddings
            chunking_strategy: Chunking strategy to use
            chunking_params: Additional parameters for the chunking strategy
            embedding_model: Embedding model to use
            storage_type: Storage type to use
            storage_params: Additional parameters for the storage adapter
            document_metadata: Additional metadata for the document

        Returns:
            Result of the chunking operation
        """
        # Validate parameters
        if not document_path:
            raise ValidationError("Document path is required")

        # Use default collection if not specified
        collection = collection_name or settings.DEFAULT_COLLECTION_NAME

        # Initialize storage adapter
        storage = await get_storage_adapter(storage_type, **(storage_params or {}))

        # Check if document exists
        if not await storage.object_exists(document_path):
            raise ResourceNotFoundError(f"Document not found: {document_path}")

        # Get document metadata
        doc_metadata = await storage.get_object_metadata(document_path)

        # Get document content for hash calculation
        document_data = await storage.get_object(document_path)
        document_hash = ProcessedDocument.calculate_hash(document_data.getvalue())

        # Check if document has already been processed
        if await self.document_tracker.is_document_processed(document_path, document_hash):
            logger.info(f"Document already processed, skipping: {document_path}")

            # Return the existing document information
            processed_doc = await self.document_tracker.get_processed_document(document_path)
            if processed_doc:
                return {
                    "document_path": processed_doc.document_path,
                    "collection_name": processed_doc.collection_name,
                    "chunk_count": processed_doc.chunk_count,
                    "chunk_ids": processed_doc.chunk_ids,
                    "chunking_strategy": processed_doc.chunking_strategy,
                    "chunking_config": processed_doc.chunking_config,
                    "already_processed": True
                }

        # Combine with provided metadata
        if document_metadata:
            doc_metadata.update(document_metadata)

        # Extract text from document
        text = await self._extract_text_from_document(storage, document_path, doc_metadata)

        # Initialize chunking strategy
        chunker = get_chunking_strategy(chunking_strategy, **(chunking_params or {}))

        # Chunk text
        chunks = chunker.chunk_text(text, doc_metadata)

        chunk_texts = [chunk["text"] for chunk in chunks]

        # Store embeddings
        ids, collection_name, count = await self.embedding_client.store_embeddings(
            texts=chunk_texts,
            collection_name=collection,
            model=embedding_model
        )

        logger.info(f"Chunked document '{document_path}' into {len(chunks)} chunks and stored in collection '{collection_name}'")

        # Create result dictionary
        result = {
            "document_path": document_path,
            "collection_name": collection_name,
            "chunk_count": count,
            "chunk_ids": ids,
            "chunking_strategy": chunker.get_strategy_name(),
            "chunking_config": chunker.get_strategy_config(),
            "already_processed": False
        }

        # Track the processed document
        processed_doc = ProcessedDocument(
            document_path=document_path,
            collection_name=collection_name,
            document_hash=document_hash,
            chunk_count=count,
            chunk_ids=ids,
            chunking_strategy=chunker.get_strategy_name(),
            chunking_config=chunker.get_strategy_config(),
            metadata=doc_metadata
        )

        await self.document_tracker.track_document(processed_doc)

        return result

    async def chunk_collection(
        self,
        collection_path: str,
        vector_collection_name: Optional[str] = None,
        chunking_strategy: Optional[str] = None,
        chunking_params: Optional[Dict[str, Any]] = None,
        embedding_model: Optional[str] = None,
        storage_type: Optional[str] = None,
        storage_params: Optional[Dict[str, Any]] = None,
        collection_metadata: Optional[Dict[str, Any]] = None,
        file_extensions: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """Chunk a collection of documents and store embeddings.

        Args:
            collection_path: Path to the collection in storage
            vector_collection_name: Name of the collection to store embeddings
            chunking_strategy: Chunking strategy to use
            chunking_params: Additional parameters for the chunking strategy
            embedding_model: Embedding model to use
            storage_type: Storage type to use
            storage_params: Additional parameters for the storage adapter
            collection_metadata: Additional metadata for the collection
            file_extensions: List of file extensions to process

        Returns:
            Result of the chunking operation
        """
        # Validate parameters
        if not collection_path:
            raise ValidationError("Collection path is required")

        # Use default collection if not specified
        vector_collection = vector_collection_name or settings.DEFAULT_COLLECTION_NAME

        # Initialize storage adapter
        storage = await get_storage_adapter(storage_type, **(storage_params or {}))

        # List objects in collection
        objects = await storage.list_objects(collection_path)

        # Filter objects by file extension if specified
        if file_extensions:
            objects = [
                obj for obj in objects
                if not obj.get("is_dir") and any(
                    obj.get("name", "").lower().endswith(ext.lower())
                    for ext in file_extensions
                )
            ]
        else:
            # Only include files, not directories
            objects = [obj for obj in objects if not obj.get("is_dir")]

        if not objects:
            logger.warning(f"No documents found in collection: {collection_path}")
            return {
                "collection_path": collection_path,
                "vector_collection_name": vector_collection,
                "document_count": 0,
                "chunk_count": 0,
                "documents": []
            }

        # Process each document
        results = []
        total_chunks = 0

        for obj in objects:
            try:
                # Create document metadata
                doc_metadata = collection_metadata.copy() if collection_metadata else {}
                doc_metadata.update({
                    "collection_path": collection_path,
                    "document_name": obj.get("name"),
                    "document_path": obj.get("path")
                })

                # Chunk document
                result = await self.chunk_document(
                    document_path=obj.get("path"),
                    collection_name=vector_collection,
                    chunking_strategy=chunking_strategy,
                    chunking_params=chunking_params,
                    embedding_model=embedding_model,
                    storage_type=storage_type,
                    storage_params=storage_params,
                    document_metadata=doc_metadata
                )

                results.append(result)
                total_chunks += result.get("chunk_count", 0)
            except Exception as e:
                logger.error(f"Failed to process document '{obj.get('path')}': {str(e)}")
                # Continue with other documents

        logger.info(f"Processed {len(results)} documents from collection '{collection_path}' with {total_chunks} total chunks")

        return {
            "collection_path": collection_path,
            "vector_collection_name": vector_collection,
            "document_count": len(results),
            "chunk_count": total_chunks,
            "documents": results
        }

    async def _extract_text_from_document(
        self,
        storage: Any,
        document_path: str,
        metadata: Dict[str, Any]
    ) -> str:
        """Extract text from a document.

        Args:
            storage: Storage adapter
            document_path: Path to the document
            metadata: Document metadata

        Returns:
            Extracted text
        """
        # Get content type
        content_type = metadata.get("content_type", "application/octet-stream")

        # Handle different content types
        if content_type.startswith("text/"):
            # Plain text
            return await storage.get_text_content(document_path)
        elif content_type == "application/pdf":
            # PDF
            return await self._extract_text_from_pdf(storage, document_path)
        elif content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]:
            # Word document
            return await self._extract_text_from_docx(storage, document_path)
        else:
            # Try to extract as text
            try:
                return await storage.get_text_content(document_path)
            except Exception as e:
                raise ChunkingError(f"Unsupported document type: {content_type}")

    async def _extract_text_from_pdf(self, storage: Any, document_path: str) -> str:
        """Extract text from a PDF document using PyMuPDF (fitz).

        Args:
            storage: Storage adapter
            document_path: Path to the document

        Returns:
            Extracted text with layout preservation
        """
        try:
            import fitz  # PyMuPDF

            # Get document data
            data = await storage.get_object(document_path)

            # Save to temporary file
            temp_path = f"/tmp/{os.path.basename(document_path)}"
            with open(temp_path, "wb") as f:
                f.write(data.getvalue())

            # Extract text with layout preservation
            text = ""
            doc = fitz.open(temp_path)

            for page_num in range(len(doc)):
                page = doc[page_num]
                # Extract text with layout preservation
                page_text = page.get_text("text")
                text += page_text + "\n\n"

            # Clean up
            os.remove(temp_path)

            return text
        except ImportError:
            raise ChunkingError("PyMuPDF is required for PDF extraction. Install with 'pip install pymupdf'")
        except Exception as e:
            raise ChunkingError(f"Failed to extract text from PDF: {str(e)}")

    async def _extract_text_from_docx(self, storage: Any, document_path: str) -> str:
        """Extract text from a Word document.

        Args:
            storage: Storage adapter
            document_path: Path to the document

        Returns:
            Extracted text
        """
        try:
            import docx

            # Get document data
            data = await storage.get_object(document_path)

            # Save to temporary file
            temp_path = f"/tmp/{os.path.basename(document_path)}"
            with open(temp_path, "wb") as f:
                f.write(data.getvalue())

            # Extract text
            doc = docx.Document(temp_path)
            text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])

            # Clean up
            os.remove(temp_path)

            return text
        except ImportError:
            raise ChunkingError("python-docx is required for DOCX extraction. Install with 'pip install python-docx'")
        except Exception as e:
            raise ChunkingError(f"Failed to extract text from DOCX: {str(e)}")

    async def list_collections(self) -> List[Dict[str, Any]]:
        """List all collections.

        Returns:
            List of collections
        """
        return await self.embedding_client.list_collections()
